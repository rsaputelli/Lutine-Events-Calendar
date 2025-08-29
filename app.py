# Streamlit Master Calendar Intake App (Final – Clean)
# Features:
# - AM/PM time selectors
# - Client/Manager dropdowns with “Other…” and auto-persist to Supabase
# - Outlook events created with showAs = "free"
# - Virtual link optional with one-click confirmation + (optional) 7‑day reminder seed
# - Optional SMTP email to Meeting Manager
# - CME/Accreditation Required email (to mkomenko@…, cc tbarrett@…)
# - Word export grouped by month

import io
from datetime import datetime, date, time, timedelta
from zoneinfo import ZoneInfo
from typing import List, Tuple

import streamlit as st
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from supabase import create_client, Client

import smtplib
from email.mime.text import MIMEText
from email.utils import formataddr

import re, html, requests
from docx.shared import RGBColor  # for red text
from urllib.parse import quote


# -----------------------------
# Config & Secrets
# -----------------------------
st.set_page_config(page_title="Lutine Master Calendar Intake", layout="wide")

# Header/logo (main page, replaces sidebar branding)
logo_col, title_col = st.columns([1, 6])
with logo_col:
    st.image("assets/lutine-logo.png", width=230)
with title_col:
    st.title("Master Calendar Intake Form")
    st.caption("Use this form to add or edit Master Calendar")


GRAPH = st.secrets.get("graph", {})
SUPA = st.secrets.get("supabase", {})
SMTP = st.secrets.get("smtp", {})  # optional: host, port, user, password, from_addr, from_name

missing = []
for k in ("tenant_id", "client_id", "client_secret", "shared_mailbox_upn"):
    if not GRAPH.get(k):
        missing.append(f"graph.{k}")
for k in ("url", "key"):
    if not SUPA.get(k):
        missing.append(f"supabase.{k}")
if missing:
    st.warning("Secrets missing: " + ", ".join(missing) + ". You can still explore the form, but submissions will be disabled.")

supabase: Client | None = None
if SUPA.get("url") and SUPA.get("key"):
    supabase = create_client(SUPA["url"], SUPA["key"])

# -----------------------------
# Helper: Time zones (US) -> Windows TZ IDs for Graph
# -----------------------------
TZ_MAP = {
    "Eastern": "Eastern Standard Time",
    "Central": "Central Standard Time",
    "Mountain": "Mountain Standard Time",
    "Pacific": "Pacific Standard Time",
    "Alaska": "Alaskan Standard Time",
    "Hawaii": "Hawaiian Standard Time",
}
IANA_MAP = {
    "Eastern": "America/New_York",
    "Central": "America/Chicago",
    "Mountain": "America/Denver",
    "Pacific": "America/Los_Angeles",
    "Alaska": "America/Anchorage",
    "Hawaii": "Pacific/Honolulu",
}


# -----------------------------
# Graph OAuth + Event Create
# -----------------------------

def get_graph_token(tenant_id: str, client_id: str, client_secret: str) -> str:
    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    data = {
        "client_id": client_id,
        "client_secret": client_secret,
        "scope": "https://graph.microsoft.com/.default",
        "grant_type": "client_credentials",
    }
    resp = requests.post(token_url, data=data, timeout=20)
    resp.raise_for_status()
    return resp.json()["access_token"]


def graph_create_event(token: str, shared_mailbox_upn: str, payload: dict) -> dict:
    url = f"https://graph.microsoft.com/v1.0/users/{shared_mailbox_upn}/calendar/events"
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    r = requests.post(url, headers=headers, json=payload, timeout=20)
    if r.status_code >= 400:
        raise RuntimeError(f"Graph error {r.status_code}: {r.text}")
    return r.json()
    
def update_outlook_event(token: str, upn: str, event_id: str, payload: dict):
    """PATCH an existing event in Outlook (used by the Edit tab)."""
    url = f"https://graph.microsoft.com/v1.0/users/{upn}/events/{event_id}"
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    r = requests.patch(url, headers=headers, json=payload, timeout=20)
    if r.status_code >= 400:
        raise RuntimeError(f"Graph PATCH {r.status_code}: {r.text}")
    return r.json()

def graph_delete_event(token: str, shared_mailbox_upn: str, outlook_event_id: str):
    """DELETE an event in Outlook. 204 = deleted; 404 = already gone."""
    url = f"https://graph.microsoft.com/v1.0/users/{shared_mailbox_upn}/events/{outlook_event_id}"
    headers = {"Authorization": f"Bearer {token}"}
    r = requests.delete(url, headers=headers, timeout=20)
    if r.status_code not in (204, 404):
        raise RuntimeError(f"Graph DELETE {r.status_code}: {r.text}")
        

def update_outlook_manager_block(outlook_event_id: str, manager_name: str, *, mailbox_upn: str, token: str) -> bool:
    """
    Ensure exactly ONE 'Meeting Manager + [App Outlook Event ID: …]' block.
    Strips any legacy p/div/span/table blocks regardless of styling, then appends one 11pt table block.
    """
    try:
        get_url = f"https://graph.microsoft.com/v1.0/users/{mailbox_upn}/events/{outlook_event_id}"
        headers = {"Authorization": f"Bearer {token}"}

        # 1) Fetch body
        r = requests.get(get_url, headers=headers, timeout=15)
        r.raise_for_status()
        body = (r.json() or {}).get("body", {}) or {}
        ctype = (body.get("contentType") or "html").lower()
        cur_html = body.get("content") or ""
        if ctype == "text" and cur_html:
            from html import escape
            cur_html = f"<pre>{escape(cur_html)}</pre>"

        # 2) Preserve an existing ID if present anywhere
        m_id = re.search(r"\[(?:App\s+)?Outlook\s+Event\s+ID:?\s*(?P<eid>[^\]]+)\]", cur_html, flags=re.I)
        preserved_id = (m_id.group("eid").strip() if m_id else "") or outlook_event_id

        # 3) Remove ANY existing Manager block variants (tables, blocks, mixed wrappers)
        # a) any <table> ... contains both phrases
        pat_table = re.compile(r"<table\b.*?>.*?Meeting\s*Manager:.*?\[.*?Outlook\s+Event\s+ID.*?\].*?</table>",
                               re.I | re.S)
        # b) any block tag (p/div/span/td) containing both phrases
        pat_block = re.compile(r"<(?P<tag>p|div|span|td)\b[^>]*>.*?Meeting\s*Manager:.*?\[.*?Outlook\s+Event\s+ID.*?\].*?</(?P=tag)>",
                               re.I | re.S)
        # c) belt-and-suspenders: inline fragment (manager … ID …) across tags within ~1500 chars
        pat_inline = re.compile(r"Meeting\s*Manager:.*?\[.*?Outlook\s+Event\s+ID.*?\].{0,50}", re.I | re.S)

        changed = True
        while changed:
            new_html = pat_table.sub("", cur_html)
            new_html = pat_block.sub("", new_html)
            new_html = pat_inline.sub("", new_html)
            changed = (new_html != cur_html)
            cur_html = new_html

        # 4) Append ONE normalized 11pt table block (Outlook-friendly)
        safe_mgr = html.escape(manager_name)
        safe_id  = html.escape(preserved_id)
        block = (
            "<table role='presentation' style='border-collapse:collapse;border-spacing:0;margin:0;padding:0;'>"
            "<tr><td style='font-family:Segoe UI, Arial, sans-serif; font-size:11pt; color:#c00000;'>"
            f"<b>Meeting Manager: {safe_mgr}</b><br><br>"
            f"<b>[App Outlook Event ID: {safe_id}]</b>"
            "</td></tr></table>"
        )
        new_html = cur_html + block

        p = requests.patch(
            get_url,
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            json={"body": {"contentType": "HTML", "content": new_html}},
            timeout=20,
        )
        p.raise_for_status()
        return True

    except Exception:
        return False

def graph_get_event(token: str, upn: str, event_id: str) -> dict:
    """GET one event by Graph ID (URL-encodes the ID)."""
    if not event_id:
        raise ValueError("event_id is required")
    url = f"https://graph.microsoft.com/v1.0/users/{upn}/events/{quote(event_id, safe='')}"
    headers = {"Authorization": f"Bearer {token}"}
    r = requests.get(url, headers=headers, timeout=20)
    r.raise_for_status()
    return r.json()


        
def graph_datetime_obj(dt_local, *, tz_windows: str) -> dict:
    """
    Convert a timezone-aware local datetime to the MS Graph event datetime object.
    Graph expects local wall time and a Windows time zone ID.
    """
    # Ensure dt_local is timezone-aware in the target local zone before formatting
    if getattr(dt_local, "tzinfo", None) is None:
        raise ValueError("dt_local must be timezone-aware")

    return {
        "dateTime": dt_local.strftime("%Y-%m-%dT%H:%M:%S"),
        "timeZone": tz_windows,  # e.g., "Eastern Standard Time"
    }
    

# -----------------------------
# Email helpers (optional)
# -----------------------------

def send_email(to_addrs, subject: str, html_body: str, cc_addrs=None):
    """Send HTML email via SMTP settings in [smtp] secrets. Returns (ok: bool, info: str)."""
    if not SMTP:
        return False, "SMTP not configured"

    # Normalize inputs to lists
    if isinstance(to_addrs, str):
        to_addrs = [to_addrs]
    if cc_addrs is None:
        cc_addrs = []
    elif isinstance(cc_addrs, str):
        cc_addrs = [cc_addrs]

    try:
        msg = MIMEText(html_body, "html")
        msg["Subject"] = subject
        from_addr = SMTP.get("from_addr", SMTP.get("user"))
        from_name = SMTP.get("from_name", "Lutine Calendar Bot")
        msg["From"] = formataddr((from_name, from_addr))
        msg["To"] = ", ".join(to_addrs)
        if cc_addrs:
            msg["Cc"] = ", ".join(cc_addrs)
        with smtplib.SMTP(SMTP.get("host"), int(SMTP.get("port", 587))) as server:
            server.starttls()
            server.login(SMTP.get("user"), SMTP.get("password"))
            server.sendmail(from_addr, to_addrs + cc_addrs, msg.as_string())
        return True, "sent"
    except Exception as e:
        return False, str(e)

# -----------------------------
# Payload builder
# -----------------------------

def build_graph_event_payload(
    subject: str,
    body_html: str,
    tz_windows: str,
    start_dt: datetime | date,
    end_dt: datetime | date,
    is_all_day: bool,
    location_str: str | None,
    set_teams: bool,
    reminder_minutes: int,
) -> dict:
    payload: dict = {
        "subject": subject,
        "isReminderOn": True,
        "reminderMinutesBeforeStart": int(reminder_minutes),
        "body": {"contentType": "HTML", "content": body_html},
        "showAs": "free",  # show as Free so multi-day all-day sits at the top
    }

    if is_all_day:
        # Date-only for all-day; Graph end is exclusive (next day)
        if isinstance(start_dt, datetime):
            start_date = start_dt.date()
        else:
            start_date = start_dt
        if isinstance(end_dt, datetime):
            end_date = end_dt.date()
        else:
            end_date = end_dt
        payload.update({
            "isAllDay": True,
            "start": {"dateTime": start_date.isoformat(), "timeZone": tz_windows},
            "end": {"dateTime": end_date.isoformat(), "timeZone": tz_windows},
        })
    else:
        payload.update({
            "start": {"dateTime": start_dt.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_windows},
            "end": {"dateTime": end_dt.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_windows},
        })

    if location_str:
        payload["location"] = {"displayName": location_str}

    if set_teams:
        payload["isOnlineMeeting"] = True
        payload["onlineMeetingProvider"] = "teamsForBusiness"

    return payload

# -----------------------------
# UI helpers – AM/PM selectors
# -----------------------------

def ampm_time_picker(label_prefix: str, default: time = time(9, 0), key_prefix: str = "") -> time:
    colh, colm, cola = st.columns([1, 1, 1])
    hour_12 = default.hour % 12
    hour_12 = 12 if hour_12 == 0 else hour_12
    ampm = "AM" if default.hour < 12 else "PM"

    h = colh.selectbox(
        f"{label_prefix} Hour", list(range(1, 13)),
        index=list(range(1, 13)).index(hour_12),
        key=f"{key_prefix}_hour"
    )
    m = colm.selectbox(
        f"{label_prefix} Min", [0, 15, 30, 45],
        index=[0, 15, 30, 45].index(default.minute if default.minute in [0, 15, 30, 45] else 0),
        key=f"{key_prefix}_min"
    )
    a = cola.selectbox(
        f"{label_prefix} AM/PM", ["AM", "PM"],
        index=["AM", "PM"].index(ampm),
        key=f"{key_prefix}_ampm"
    )

    hh = (h % 12) + (12 if a == "PM" and h != 12 else 0)
    if a == "AM" and h == 12:
        hh = 0
    return time(hh, m)


# -----------------------------
# Load dropdown data (clients & managers)
# -----------------------------

def load_clients() -> List[str]:
    try:
        if supabase is None:
            return []
        res = supabase.table("clients").select("name").order("name").execute()
        return [r["name"] for r in (res.data or [])]
    except Exception:
        return []


def load_managers() -> List[Tuple[str, str]]:
    try:
        if supabase is None:
            return []
        res = supabase.table("meeting_managers").select("name,email").order("name").execute()
        return [(r["name"], r.get("email", "")) for r in (res.data or [])]
    except Exception:
        return []
        
# ---- Graph delta bookmark helpers (Supabase) ----
def get_delta_link(scope: str = "default") -> str | None:
    if supabase is None:
        return None
    res = supabase.table("graph_state").select("delta_link").eq("scope", scope).limit(1).execute()
    rows = res.data or []
    return rows[0]["delta_link"] if rows and rows[0].get("delta_link") else None

def save_delta_link(delta_link: str, scope: str = "default") -> None:
    if supabase is None:
        return
    supabase.table("graph_state").upsert({
        "scope": scope,
        "delta_link": delta_link,
        "last_synced": datetime.utcnow().isoformat()
    }, on_conflict="scope").execute()
    
# ---------- Graph GET single event ----------
def graph_get_event(token: str, shared_mailbox_upn: str, event_id: str) -> dict:
    url = f"https://graph.microsoft.com/v1.0/users/{shared_mailbox_upn}/events/{event_id}"
    headers = {"Authorization": f"Bearer {token}"}
    r = requests.get(url, headers=headers, timeout=20)
    if r.status_code >= 400:
        raise RuntimeError(f"Graph getEvent {r.status_code}: {r.text}")
    return r.json()

# ---------- Graph delta (calendarView) ----------
def graph_delta_events(token: str, shared_mailbox_upn: str, start_iso: str | None, end_iso: str | None, delta_link: str | None = None):
    """
    If delta_link is provided, call it directly (it already includes query params).
    Otherwise, call calendarView/delta with a UTC window using requests' params= to ensure proper URL encoding.
    Yields page dicts; each page may contain '@odata.nextLink' or '@odata.deltaLink'.
    """
    headers = {"Authorization": f"Bearer {token}"}

    if delta_link:
        next_url = delta_link
        while next_url:
            r = requests.get(next_url, headers=headers, timeout=30)
            if r.status_code >= 400:
                raise RuntimeError(f"Graph delta {r.status_code}: {r.text}")
            page = r.json()
            yield page
            next_url = page.get("@odata.nextLink")
        return

    # First-time windowed delta (encode params properly)
    base = f"https://graph.microsoft.com/v1.0/users/{shared_mailbox_upn}/calendarView/delta"
    params = {}
    if start_iso: params["startDateTime"] = start_iso
    if end_iso:   params["endDateTime"] = end_iso

    next_link = None
    while True:
        if next_link:
            r = requests.get(next_link, headers=headers, timeout=30)
        else:
            r = requests.get(base, headers=headers, params=params, timeout=30)
        if r.status_code >= 400:
            raise RuntimeError(f"Graph delta {r.status_code}: {r.text}")
        page = r.json()
        yield page
        next_link = page.get("@odata.nextLink")
        if not next_link:
            break
  
# ---------- Outlook → App field mapping helpers ----------
from zoneinfo import ZoneInfo
from datetime import datetime

def _parse_graph_dt_to_utc(dt_obj: dict) -> str | None:
    """
    Convert Graph dateTime dict -> UTC ISO string.
    Accepts values like:
      {"dateTime":"2025-09-01T09:00:00.0000000","timeZone":"Eastern Standard Time"}
    Logic:
      - If string has a 'Z' or offset, we trust it and convert to UTC.
      - If naive, treat as UTC (safe fallback for our use case).
    """
    if not dt_obj or not dt_obj.get("dateTime"):
        return None
    dt_raw = dt_obj["dateTime"]
    try:
        dt = datetime.fromisoformat(dt_raw.replace("Z", "+00:00"))
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=ZoneInfo("UTC"))
        return dt.astimezone(ZoneInfo("UTC")).isoformat()
    except Exception:
        return None

def map_graph_event_to_row_updates(g: dict) -> dict:
    """
    Return ONLY the Outlook-owned fields you want to overwrite in 'events'.
    Internal app fields (deliverables, accreditation, manager, etc.) are untouched.
    """
    updates = {}

    # Subject (optional—enable if you want Outlook title to win)
    subj = g.get("subject")
    if subj:
        updates["subject"] = subj

    # All-day
    updates["is_all_day"] = bool(g.get("isAllDay"))

    # Start/End (UTC ISO)
    start_utc = _parse_graph_dt_to_utc(g.get("start"))
    end_utc   = _parse_graph_dt_to_utc(g.get("end"))
    if start_utc: updates["start_dt_utc"] = start_utc
    if end_utc:   updates["end_dt_utc"] = end_utc

    # Location (only physical display name)
    loc = (g.get("location") or {}).get("displayName") or ""
    if loc:
        updates["location"] = loc

    # Teams/online meeting URL → virtual_link (only fill/refresh the link)
    om = g.get("onlineMeeting")
    if isinstance(om, dict):
        join_url = om.get("joinUrl")
        if join_url:
            updates["virtual_link"] = join_url

    return updates
      

# -----------------------------
# Formatting helper for emails
# -----------------------------

def fmt_event_info(subject: str, start_dt_et: datetime, end_dt_et: datetime,
                   is_all_day: bool, tz_label: str, client_value: str,
                   event_type: str, location: str, vp_label: str,
                   virtual_link: str, manager_name: str, manager_email: str) -> str:
    if is_all_day:
        # For all-day, show date(s) only (remember Graph end is exclusive)
        if (end_dt_et - timedelta(days=1)).date() == start_dt_et.date():
            when = start_dt_et.strftime("%B %d, %Y")
        else:
            when = f"{start_dt_et.strftime('%B %d, %Y')} – {(end_dt_et - timedelta(days=1)).strftime('%B %d, %Y')}"
    else:
        when = f"{start_dt_et.strftime('%B %d, %Y, %I:%M %p')} – {end_dt_et.strftime('%I:%M %p')} {tz_label}"

    parts = [
        f"<p><b>Event:</b> {subject}</p>",
        f"<p><b>When:</b> {when}</p>",
    ]

    if client_value:
        parts.append(f"<p><b>Client:</b> {client_value}</p>")
    if event_type == "In-person" and location:
        parts.append(f"<p><b>Location:</b> {location}</p>")
    elif event_type == "Virtual":
        parts.append(f"<p><b>Virtual:</b> {vp_label}{(' – ' + virtual_link) if virtual_link else ''}</p>")
    if manager_name or manager_email:
        parts.append(f"<p><b>Meeting Manager:</b> {manager_name} {('<' + manager_email + '>') if manager_email else ''}</p>")

    return "\n".join(parts)


# -----------------------------
# UI – Create & Edit Tabs
# -----------------------------
tab_create, tab_edit = st.tabs(["Create", "Edit"])



# ==========
# CREATE TAB
# ==========
with tab_create:
    st.subheader("Create Event & Post to Master Calendar")

    # -- Persist all-day flag outside the form so conditional fields update immediately
    if "is_all_day" not in st.session_state:
        st.session_state["is_all_day"] = False

    event_type = st.selectbox("Event Type", ["In-person", "Virtual"], index=0, key="create_event_type")
    is_all_day = st.checkbox("All-Day Event", value=st.session_state["is_all_day"], key="create_is_all_day")
    st.session_state["is_all_day"] = is_all_day

    prev_type = st.session_state.get("prev_event_type_create")
    if prev_type != event_type:
        st.session_state["confirm_no_link"] = False
    st.session_state["prev_event_type_create"] = event_type

    # --- Reminder controls OUTSIDE the form so they re-render live ---
    rem_col1, rem_col2 = st.columns([1, 2])
    if "rem_mode" not in st.session_state:
        st.session_state["rem_mode"] = "Minutes before start (Outlook)"

    st.session_state["rem_mode"] = rem_col1.selectbox(
        "Reminder Type",
        ["Minutes before start (Outlook)", "Days before start (Outlook)", "On date/time (Email via app)"],
        index=["Minutes before start (Outlook)", "Days before start (Outlook)", "On date/time (Email via app)"]
              .index(st.session_state["rem_mode"]),
        key="create_rem_mode_live",
    )

    # Live companion input for reminder
    reminder_minutes = st.session_state.get("reminder_minutes", 30)
    reminder_days = st.session_state.get("reminder_days", 1)
    reminder_datetime_local = st.session_state.get("reminder_datetime_local")

    if st.session_state["rem_mode"].startswith("Minutes"):
        reminder_minutes = rem_col2.number_input(
            "Minutes before start", min_value=0, max_value=10080, value=int(reminder_minutes), key="create_rem_mins_live"
        )
        st.session_state["reminder_minutes"] = reminder_minutes

    elif st.session_state["rem_mode"].startswith("Days"):
        reminder_days = rem_col2.number_input(
            "Days before start", min_value=0, max_value=365, value=int(reminder_days), key="create_rem_days_live"
        )
        st.session_state["reminder_days"] = reminder_days

    else:
        default_dt = datetime.combine(date.today(), time(9, 0))
        reminder_datetime_local = rem_col2.datetime_input(
            "Reminder date & time", value=reminder_datetime_local or default_dt, key="create_rem_dt_live"
        )
        st.session_state["reminder_datetime_local"] = reminder_datetime_local

    # -----------------
    # Create Event Form
    # -----------------
    with st.form("event_form_create"):
        col1, col2, col3 = st.columns(3)
        subject = col1.text_input("Event Title *", "", key="create_subject")

        # Client dropdown with Other
        client_options = load_clients()
        client_sel = col2.selectbox("Client", client_options + ["Other…"],
                                    index=(0 if client_options else 0), key="create_client_sel")
        client_other = ""
        if client_sel == "Other…":
            client_other = col2.text_input("Enter new client name", "", key="create_client_other")
        client_value = client_other if client_sel == "Other…" else (client_sel or "")

        tz_choice = col3.selectbox("Time Zone", list(TZ_MAP.keys()), index=0, key="create_tz")

        # Dates (always in the form)
        col4, col5 = st.columns(2)
        start_date = col4.date_input("Start Date", value=date.today(), key="create_start_date")
        end_date   = col5.date_input("End Date",   value=date.today(), key="create_end_date")

        # Times (only if not all-day) — allow typing with time_input
        if not is_all_day:
            start_time = st.time_input("Start Time", value=time(9, 0), key="create_start_time",
                                       step=timedelta(minutes=5))
            end_time   = st.time_input("End Time",   value=time(10, 0), key="create_end_time",
                                       step=timedelta(minutes=5))
        else:
            start_time = time(0, 0)
            end_time   = time(0, 0)

        # Event-type-specific fields
        location = ""
        virtual_provider = None
        virtual_link = None
        if event_type == "In-person":
            location = st.text_input("Location (City, Venue, etc.) *", "", key="create_location")
        else:
            virtual_provider_label = st.selectbox("Virtual Platform", ["Teams", "Zoom", "Other"],
                                                  index=0, key="create_vp_label")
            PROVIDER_MAP = {"Teams": "teams", "Zoom": "zoom", "Other": "other"}
            virtual_provider = PROVIDER_MAP.get(virtual_provider_label, "other")
            virtual_link = st.text_input("Virtual Meeting Link (optional)", "", key="create_vlink")

        st.markdown("---")
        accreditation_required = st.selectbox("CME/Accreditation Required?", ["No", "Yes"],
                                              index=0, key="create_acc") == "Yes"
        st.markdown("---")

        # Meeting Manager dropdown with Other
        st.markdown("**Meeting Manager (internal only):**")
        managers = load_managers()  # list of (name,email)
        manager_labels = [f"{n} <{e}>" if e else n for n, e in managers]
        manager_sel = st.selectbox("Choose manager", manager_labels + ["Other…"],
                                   index=(0 if managers else 0), key="create_mm_sel")
        manager_name = ""
        manager_email = ""
        if manager_sel == "Other…":
            mm_col1, mm_col2 = st.columns(2)
            manager_name = mm_col1.text_input("Name *", "", key="create_mm_name")
            manager_email = mm_col2.text_input("Email *", "", key="create_mm_email")
        else:
            idx = manager_labels.index(manager_sel) if manager_sel in manager_labels else -1
            if idx >= 0:
                manager_name, manager_email = managers[idx]

        # Optional Notes (included in Outlook body)
        notes = st.text_area("Notes (included in Outlook event body)", key="create_notes")

        # Confirmation for blank virtual link
        if event_type == "Virtual" and (not virtual_link):
            if not st.session_state.get("confirm_no_link"):
                st.warning("Virtual link is blank. Click again to confirm creating without a link. "
                           "We'll remind you every 7 days until a link is added.")

        submitted = st.form_submit_button("Create Event")

    # ----- Create submission handling -----
    if submitted:
        if event_type == "Virtual" and (not virtual_link) and not st.session_state.get("confirm_no_link"):
            st.session_state["confirm_no_link"] = True
            st.stop()

        errs = []
        if not subject:
            errs.append("Event Title is required.")
        if event_type == "In-person" and not location:
            errs.append("Location is required for in-person events.")
        if manager_sel == "Other…" and (not manager_name or not manager_email):
            errs.append("Meeting Manager name and email are required.")
        if errs:
            st.error("\n".join(errs))
            st.stop()

        # Build start/end in local zone
        iana = IANA_MAP[tz_choice]
        tz = ZoneInfo(iana)
        if is_all_day:
            start_dt_local = datetime.combine(start_date, time(0, 0)).replace(tzinfo=tz)
            end_base = max(end_date, start_date)
            end_dt_local = datetime.combine(end_base + timedelta(days=1), time(0, 0)).replace(tzinfo=tz)
        else:
            start_dt_local = datetime.combine(start_date, start_time).replace(tzinfo=tz)
            end_dt_local = datetime.combine(end_date, end_time).replace(tzinfo=tz)
        if not is_all_day and end_dt_local <= start_dt_local:
            st.error("End date/time must be after the start date/time. Please adjust and resubmit.")
            st.stop()
        start_dt_utc = start_dt_local.astimezone(ZoneInfo("UTC"))
        end_dt_utc = end_dt_local.astimezone(ZoneInfo("UTC"))

        # Persist “Other…” choices
        if supabase and client_value and client_sel == "Other…":
            try:
                supabase.table("clients").insert({"name": client_value}).execute()
            except Exception:
                pass
        if supabase and manager_sel == "Other…" and manager_name and manager_email:
            try:
                supabase.table("meeting_managers").insert({"name": manager_name, "email": manager_email}).execute()
            except Exception:
                pass

        # --- Minimal Outlook body: Client + Accreditation + Manager (11pt) ---
        import html  # at top of file if not already

        # Line 1: Client (optional)
        client_line = f"<p><b>Client:</b> {html.escape(client_value)}</p>" if client_value else ""

        # Line 2: Accreditation (always show Yes/No)
        acc_flag = "Yes" if accreditation_required else "No"
        acc_line = f"<p><b>Accreditation:</b> {acc_flag}</p>"

        # Line 3: Manager block (robust for Outlook using 1-cell table, 11pt)
        safe_mgr = html.escape(manager_name or "")
        manager_block = (
            "<table role='presentation' style='border-collapse:collapse;border-spacing:0;margin:0;padding:0;'>"
            "<tr><td style='font-family:Segoe UI, Arial, sans-serif; font-size:11pt; color:#c00000;'>"
            f"<b>Meeting Manager: {safe_mgr}</b><br><br>"
            f"<b>[App Outlook Event ID will sync here]</b>"
            "</td></tr></table>"
        )

        combined_body = client_line + acc_line + manager_block


        # Graph payload + reminder minutes (read from session)
        tz_windows = TZ_MAP[tz_choice]
        set_teams = (event_type == "Virtual" and virtual_provider == "teams")
        location_str = location if event_type == "In-person" else (virtual_link if virtual_provider == "zoom" else None)

        rem_mode = st.session_state["rem_mode"]
        if rem_mode.startswith("Minutes"):
            rem_minutes_for_graph = int(st.session_state.get("reminder_minutes", 30))
        elif rem_mode.startswith("Days"):
            rem_minutes_for_graph = int(st.session_state.get("reminder_days", 1)) * 1440
        else:
            rem_minutes_for_graph = 0  # date-certain handled via notifications below

        rem_minutes_for_graph = max(0, min(rem_minutes_for_graph, 525600))  # <= 365 days

        payload = build_graph_event_payload(
            subject=subject,
            body_html=combined_body,   # IMPORTANT: send styled body on create
            tz_windows=tz_windows,
            start_dt=start_dt_local if not is_all_day else start_date,
            end_dt=end_dt_local if not is_all_day else end_date,
            is_all_day=is_all_day,
            location_str=location_str,
            set_teams=set_teams,
            reminder_minutes=rem_minutes_for_graph
        )

        # ---- Create in Outlook (styled body) + PATCH placeholder -> real ID ----
        outlook_event_id = None
        patched_body_for_db = combined_body  # default; replaced if PATCH succeeds
        try:
            token = get_graph_token(GRAPH["tenant_id"], GRAPH["client_id"], GRAPH["client_secret"])
            created = graph_create_event(token, GRAPH["shared_mailbox_upn"], payload)
            outlook_event_id = (created or {}).get("id")

            if outlook_event_id:
                patch_url = f"https://graph.microsoft.com/v1.0/users/{GRAPH['shared_mailbox_upn']}/events/{outlook_event_id}"
                headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

                # swap the placeholder with the real ID, preserving the style
                patched_body_for_db = combined_body.replace(
                    "[App Outlook Event ID will sync here]",
                    f"[App Outlook Event ID: {outlook_event_id}]"
                )

                patch_payload = {"body": {"contentType": "HTML", "content": patched_body_for_db}}
                patch_resp = requests.patch(patch_url, headers=headers, json=patch_payload, timeout=15)
                patch_resp.raise_for_status()

        except Exception as e:
            st.warning(f"Outlook create/patch issue: {e}")

        # ---------- Persist in Supabase ----------
        inserted_event_id = None
        try:
            if supabase is None:
                raise RuntimeError("Supabase not configured.")
            row = {
                "subject": subject,
                "client": client_value or None,
                "start_dt_utc": start_dt_utc.isoformat(),
                "end_dt_utc": end_dt_utc.isoformat(),
                "timezone_display": iana,
                "is_all_day": is_all_day,
                "location": location or None,
                "event_type": ("virtual" if event_type == "Virtual" else "in_person"),
                "virtual_provider": (virtual_provider or None),
                "virtual_link": (virtual_link or None),
                "meeting_manager_name": manager_name,
                "meeting_manager_email": manager_email,
                "reminder_minutes": int(rem_minutes_for_graph),
                "outlook_event_id": outlook_event_id,
                "accreditation_required": bool(accreditation_required),
                "created_at": datetime.utcnow().isoformat(),
                "outlook_body_html": patched_body_for_db,  # final body we sent to Outlook
            }
            res_insert = supabase.table("events").insert(row).execute()
            if res_insert.data and len(res_insert.data) > 0:
                inserted_event_id = res_insert.data[0].get("id")

            # Date-certain reminder notification (email via app)
            if rem_mode.startswith("On date/time") and st.session_state.get("reminder_datetime_local"):
                try:
                    notify_utc = st.session_state["reminder_datetime_local"].replace(
                        tzinfo=ZoneInfo(IANA_MAP[tz_choice])
                    ).astimezone(ZoneInfo("UTC"))
                    supabase.table("notifications").insert({
                        "event_id": inserted_event_id,
                        "type": "custom_email",
                        "notify_at_utc": notify_utc.isoformat(),
                        "channel": "email",
                        "payload": {
                            "to": manager_email,
                            "subject": f"Reminder: {subject}",
                            "body": f"Reminder for {subject} ({client_value})"
                        }
                    }).execute()
                except Exception:
                    pass

            # Missing link 7-day seed
            if event_type == "Virtual" and not virtual_link:
                try:
                    notify_at = datetime.utcnow() + timedelta(days=7)
                    supabase.table("notifications").insert({
                        "event_id": inserted_event_id,
                        "type": "missing_link",
                        "notify_at_utc": notify_at.isoformat(),
                        "channel": "email",
                    }).execute()
                except Exception:
                    pass

        except Exception as e:
            st.error(f"Supabase insert failed: {e}")
        else:
            st.success("Event created and saved successfully.")

            # Optional manager email (uses your existing send_email)
            if manager_email:
                ok_mgr, info_mgr = send_email(
                    [manager_email],
                    f"You are the Meeting Manager for '{subject}'",
                    f"<p>Hello {manager_name},</p><p>You have been added as the Meeting Manager for <b>{subject}</b>.</p>"
                )
                if ok_mgr:
                    st.info("Notification email sent to Meeting Manager.")

            # Accreditation email
            if accreditation_required:
                start_et = start_dt_utc.astimezone(ZoneInfo("America/New_York"))
                end_et = end_dt_utc.astimezone(ZoneInfo("America/New_York"))
                vp_label = {"teams": "Teams", "zoom": "Zoom", "other": "Virtual"}.get(virtual_provider, "Virtual")
                info_html = fmt_event_info(subject, start_et, end_et, is_all_day, tz_choice, client_value,
                                           event_type, location, vp_label, virtual_link, manager_name, manager_email)
                ok_acc, info_acc = send_email(
                    to_addrs=["mkomenko@lutinemanagement.com"],
                    cc_addrs=["tbarrett@lutinemanagement.com"],
                    subject="Accreditation Request",
                    html_body=("<p>An event has been created that requires accreditation.</p>" + info_html)
                )
                if ok_acc:
                    st.info("Accreditation request email sent.")


# ========
# EDIT TAB
# ========
def update_outlook_event(token: str, upn: str, event_id: str, payload: dict):
    url = f"https://graph.microsoft.com/v1.0/users/{upn}/events/{event_id}"
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    r = requests.patch(url, headers=headers, json=payload, timeout=20)
    if r.status_code >= 400:
        raise RuntimeError(f"Graph PATCH {r.status_code}: {r.text}")
    return r.json()

def upsert_custom_reminder(supabase_client: Client, event_id: str, notify_at_utc: str, to_email: str, subject_line: str, body_html: str):
    existing = supabase_client.table("notifications").select("id").eq("event_id", event_id).eq("type", "custom_email").execute()
    payload = {"to": to_email, "subject": subject_line, "body": body_html}
    if existing.data:
        nid = existing.data[0]["id"]
        supabase_client.table("notifications").update({"notify_at_utc": notify_at_utc, "payload": payload}).eq("id", nid).execute()
    else:
        supabase_client.table("notifications").insert({
            "event_id": event_id, "type": "custom_email", "channel": "email",
            "notify_at_utc": notify_at_utc, "payload": payload
        }).execute()

def delete_missing_link_reminders(supabase_client: Client, event_id: str):
    supabase_client.table("notifications").delete().eq("event_id", event_id).eq("type", "missing_link").execute()

with tab_edit:
    st.subheader("Edit Existing Event")

    # Filters
    filt_cols = st.columns([2,2,2,2])
    f_start = filt_cols[0].date_input("From", value=date.today().replace(day=1), key="edit_from")
    f_end   = filt_cols[1].date_input("To", value=date.today() + timedelta(days=30), key="edit_to")
    clients = load_clients()
    f_client = filt_cols[2].selectbox("Client filter (optional)", ["(all)"] + clients, index=0, key="edit_client")

    # Load events
    edit_events = []
    if supabase is None:
        st.info("Supabase not configured.")
    else:
        try:
            q = supabase.table("events").select("*").gte("start_dt_utc", datetime.combine(f_start, time(0,0)).isoformat()).lte("start_dt_utc", datetime.combine(f_end, time(23,59)).isoformat()).order("start_dt_utc", desc=False)
            if f_client and f_client != "(all)":
                q = q.eq("client", f_client)
            res = q.execute()
            edit_events = res.data or []
        except Exception as e:
            st.error(f"Failed to load events: {e}")

    if not edit_events:
        st.caption("No events found for the selected filters.")
        st.stop()

    # Select event to edit
    options = [f"{ev['subject']} • {ev.get('client') or ''} • {ev['start_dt_utc'][:16]}" for ev in edit_events]
    sel_idx = st.selectbox("Pick an event to edit", list(range(len(options))), format_func=lambda i: options[i], key="edit_pick")
    ev = edit_events[sel_idx]
    # --- Seed session_state from the selected event (only when selection changes) ---
    selected_id = ev["id"]
    prev_selected = st.session_state.get("edit_selected_id")
    if prev_selected != selected_id:
        st.session_state["edit_selected_id"] = selected_id

        # Convert stored UTC to the event's local tz for seeding
        iana_e = ev.get("timezone_display") or "America/New_York"
        tz_e = ZoneInfo(iana_e)
        start_e_utc = datetime.fromisoformat(ev["start_dt_utc"].replace("Z", "+00:00"))
        end_e_utc   = datetime.fromisoformat(ev["end_dt_utc"].replace("Z", "+00:00"))
        start_e = start_e_utc.astimezone(tz_e)
        end_e   = end_e_utc.astimezone(tz_e)

        # Find the matching Windows tz key you use in the UI (reverse IANA->label)
        # This assumes you have IANA_MAP and TZ_MAP (Windows label -> Windows name) already.
        # If you have a map from tz label -> IANA in IANA_MAP, reverse it:
        tz_label_from_iana = next((label for label, iana in IANA_MAP.items() if iana == iana_e), list(TZ_MAP.keys())[0])
        st.session_state["edit_tz_choice"] = tz_label_from_iana

        st.session_state["edit_is_all_day"] = bool(ev.get("is_all_day"))

        # Seed dates/times
        st.session_state["edit_start_date"] = start_e.date()
        st.session_state["edit_end_date"]   = (end_e - timedelta(days=1)).date() if ev.get("is_all_day") else end_e.date()
        st.session_state["edit_start_time"] = start_e.time().replace(second=0, microsecond=0)
        st.session_state["edit_end_time"]   = (time(0,0) if ev.get("is_all_day") else end_e.time().replace(second=0, microsecond=0))

        # Seed reminder widgets sensibly
        mins = int(ev.get("reminder_minutes") or 0)
        if mins >= 1440:
            st.session_state["edit_rem_mode"] = "Days before start (Outlook)"
            st.session_state["edit_reminder_days"] = max(1, mins // 1440)
        elif mins > 0:
            st.session_state["edit_rem_mode"] = "Minutes before start (Outlook)"
            st.session_state["edit_reminder_minutes"] = mins
        else:
            st.session_state["edit_rem_mode"] = "On date/time (Email via app)"
            st.session_state["edit_reminder_datetime_local"] = datetime.combine(date.today(), time(9,0))
    
    st.markdown("### Danger Zone")
    c1, c2 = st.columns([1, 3])
    confirm_del = c1.checkbox("Yes, delete this event", key="confirm_delete_ev")
    if c2.button("Delete Event", type="secondary", disabled=not confirm_del):
        try:
            # 1) Delete from Outlook first (if present)
            if ev.get("outlook_event_id") and not missing:
                tok = get_graph_token(GRAPH["tenant_id"], GRAPH["client_id"], GRAPH["client_secret"])
                try:
                    graph_delete_event(tok, GRAPH["shared_mailbox_upn"], ev["outlook_event_id"])
                except Exception as e:
                    # Don’t block DB cleanup if Outlook delete had a hiccup
                    st.warning(f"Outlook delete issue (continuing): {e}")

            # 2) Delete from Supabase; notifications will cascade via FK
            supabase.table("events").delete().eq("id", ev["id"]).execute()

            st.success("Event deleted.")
            st.stop()  # stop to refresh UI cleanly
        except Exception as e:
            st.error(f"Delete failed: {e}")
  

    # Prefill fields
    subject_e = st.text_input("Event Title *", ev.get("subject") or "", key="edit_subject")
    tz_label_e = "Eastern"  # default for display; use stored timezone_display to infer
    iana_e = ev.get("timezone_display") or "America/New_York"
    tz_e = ZoneInfo(iana_e)

    # Derive local datetimes from UTC
    start_e_utc = datetime.fromisoformat(ev["start_dt_utc"].replace("Z", "+00:00"))
    end_e_utc   = datetime.fromisoformat(ev["end_dt_utc"].replace("Z", "+00:00"))
    start_e = start_e_utc.astimezone(tz_e)
    end_e   = end_e_utc.astimezone(tz_e)
    is_all_day_e = bool(ev.get("is_all_day"))

    # Top controls
    top_cols = st.columns(3)
    tz_choice_e = top_cols[0].selectbox(
        "Time Zone",
        list(TZ_MAP.keys()),
        index=list(TZ_MAP.keys()).index(st.session_state["edit_tz_choice"]),
        key="edit_tz_choice"
    )
    is_all_day_e = top_cols[1].checkbox("All-Day Event", key="edit_is_all_day")

    # Dates
    colD1, colD2 = st.columns(2)
    start_date_e = colD1.date_input("Start Date", key="edit_start_date")
    end_date_e   = colD2.date_input("End Date",   key="edit_end_date")

    # Times (typing allowed, not dropdown-only)
    if not is_all_day_e:
        start_time_e = st.time_input("Start Time", key="edit_start_time", step=timedelta(minutes=5))
        end_time_e   = st.time_input("End Time",   key="edit_end_time",   step=timedelta(minutes=5))
    else:
        start_time_e = time(0, 0)
        end_time_e   = time(0, 0)

    # Event type & location/virtual
    event_type_e = st.selectbox("Event Type", ["In-person", "Virtual"], index=(0 if ev.get("event_type") == "in_person" else 1), key="edit_event_type")
    location_e = ""
    virtual_provider_e = None
    virtual_link_e = None
    if event_type_e == "In-person":
        location_e = st.text_input("Location (City, Venue, etc.) *", ev.get("location") or "", key="edit_location")
    else:
        default_vp = {"teams": "Teams", "zoom": "Zoom"}.get((ev.get("virtual_provider") or "other").lower(), "Other")
        vp_label_e = st.selectbox("Virtual Platform", ["Teams", "Zoom", "Other"], index=["Teams","Zoom","Other"].index(default_vp), key="edit_vp_label")
        PROVIDER_MAP = {"Teams": "teams", "Zoom": "zoom", "Other": "other"}
        virtual_provider_e = PROVIDER_MAP.get(vp_label_e, "other")
        virtual_link_e = st.text_input("Virtual Meeting Link (optional)", ev.get("virtual_link") or "", key="edit_vlink")

    # Manager & accreditation
    st.markdown("---")
    accreditation_required_e = st.selectbox("CME/Accreditation Required?", ["No", "Yes"], index=(1 if ev.get("accreditation_required") else 0), key="edit_acc") == "Yes"
    st.markdown("---")
    managers_e = load_managers()
    manager_labels_e = [f"{n} <{e}>" if e else n for n, e in managers_e]
    # preselect matching manager if possible
    default_label = f"{ev.get('meeting_manager_name') or ''} <{ev.get('meeting_manager_email') or ''}>".strip()
    try_idx = manager_labels_e.index(default_label) if default_label in manager_labels_e else 0
    manager_sel_e = st.selectbox("Choose manager", manager_labels_e + ["Other…"], index=(try_idx if manager_labels_e else 0), key="edit_mm_sel")
    manager_name_e = ev.get("meeting_manager_name") or ""
    manager_email_e = ev.get("meeting_manager_email") or ""
    if manager_sel_e == "Other…":
        mm2c1, mm2c2 = st.columns(2)
        manager_name_e = mm2c1.text_input("Name *", manager_name_e, key="edit_mm_name")
        manager_email_e = mm2c2.text_input("Email *", manager_email_e, key="edit_mm_email")
    else:
        idx2 = manager_labels_e.index(manager_sel_e) if manager_sel_e in manager_labels_e else -1
        if idx2 >= 0:
            manager_name_e, manager_email_e = managers_e[idx2]

    # Reminder modes (edit)
    rem2c1, rem2c2 = st.columns([1, 2])

    # Initialize once from stored minutes
    if "edit_rem_mode" not in st.session_state:
        stored_mins = int(ev.get("reminder_minutes") or 0)
        st.session_state["edit_rem_mode"] = (
            "Minutes before start (Outlook)" if stored_mins and stored_mins < 1440 else
            ("Days before start (Outlook)" if stored_mins >= 1440 else "On date/time (Email via app)")
        )
    if "edit_reminder_minutes" not in st.session_state:
        st.session_state["edit_reminder_minutes"] = int(ev.get("reminder_minutes") or 30)
    if "edit_reminder_days" not in st.session_state:
        base = int(ev.get("reminder_minutes") or 0)
        st.session_state["edit_reminder_days"] = max(1, base // 1440) if base >= 1440 else 1
    if "edit_reminder_datetime_local" not in st.session_state:
        st.session_state["edit_reminder_datetime_local"] = datetime.combine(date.today(), time(9,0))

    st.session_state["edit_rem_mode"] = rem2c1.selectbox(
        "Reminder Type",
        ["Minutes before start (Outlook)", "Days before start (Outlook)", "On date/time (Email via app)"],
        index=["Minutes before start (Outlook)", "Days before start (Outlook)", "On date/time (Email via app)"]
              .index(st.session_state["edit_rem_mode"]),
        key="edit_rem_mode_live"
    )

    if st.session_state["edit_rem_mode"].startswith("Minutes"):
        st.session_state["edit_reminder_minutes"] = rem2c2.number_input(
            "Minutes before start", min_value=0, max_value=10080,
            value=int(st.session_state["edit_reminder_minutes"]),
            key="edit_rem_mins_live"
        )
    elif st.session_state["edit_rem_mode"].startswith("Days"):
        st.session_state["edit_reminder_days"] = rem2c2.number_input(
            "Days before start", min_value=1, max_value=365,
            value=int(st.session_state["edit_reminder_days"]),
            key="edit_rem_days_live"
        )
    else:
        st.session_state["edit_reminder_datetime_local"] = rem2c2.datetime_input(
            "Reminder date & time",
            value=st.session_state["edit_reminder_datetime_local"],
            key="edit_rem_dt_live"
        )


    # 👇 ADD THIS BLOCK *HERE* (just before the Save button)
    st.text_input(
        "Outlook Event ID",
        value=ev.get("outlook_event_id") or "",
        disabled=True,
        key="edit_outlook_event_id_ro"
    )

    if st.button("Save Changes", type="primary"):
        # ---- Validate ----
        errs = []
        if not subject_e:
            errs.append("Event Title is required.")
        if event_type_e == "In-person" and not location_e:
            errs.append("Location is required for in-person events.")
        if not (manager_name_e and manager_email_e):
            errs.append("Meeting Manager name and email are required.")
        if errs:
            st.error("\n".join(errs))
            st.stop()

        with st.spinner("Updating event…"):
            try:
                # ---- Build new local datetimes ----
                tz_choice_lbl = tz_choice_e
                iana_new = IANA_MAP[tz_choice_lbl]
                tz_new = ZoneInfo(iana_new)

                if is_all_day_e:
                    start_local_new = datetime.combine(start_date_e, time(0, 0)).replace(tzinfo=tz_new)
                    end_base_new = max(end_date_e, start_date_e)
                    end_local_new = datetime.combine(end_base_new + timedelta(days=1), time(0, 0)).replace(tzinfo=tz_new)
                else:
                    start_local_new = datetime.combine(start_date_e, start_time_e).replace(tzinfo=tz_new)
                    end_local_new = datetime.combine(end_date_e, end_time_e).replace(tzinfo=tz_new)

                if not is_all_day_e and end_local_new <= start_local_new:
                    st.error("End date/time must be after start date/time.")
                    st.stop()

                start_utc_new = start_local_new.astimezone(ZoneInfo("UTC"))
                end_utc_new   = end_local_new.astimezone(ZoneInfo("UTC"))

                # ---- Reminder minutes for Outlook ----
                rem_mode_e = st.session_state["edit_rem_mode"]
                if rem_mode_e.startswith("Minutes"):
                    rem_minutes_for_graph_e = int(st.session_state.get("edit_reminder_minutes", 30))
                elif rem_mode_e.startswith("Days"):
                    rem_minutes_for_graph_e = int(st.session_state.get("edit_reminder_days", 1)) * 1440
                else:
                    rem_minutes_for_graph_e = 0  # date-certain handled below
                rem_minutes_for_graph_e = max(0, min(rem_minutes_for_graph_e, 525600))

            # ---- PATCH Outlook core fields (subject/time/location/reminder); DO NOT send body ----
            if ev.get("outlook_event_id"):
                if missing:
                    raise RuntimeError("Missing Graph secrets for update.")
                token = get_graph_token(GRAPH["tenant_id"], GRAPH["client_id"], GRAPH["client_secret"])

                tz_windows_e = TZ_MAP[tz_choice_e]
                patch_payload = {
                    "subject": subject_e,
                    "isAllDay": bool(is_all_day_e),
                    "start": graph_datetime_obj(start_local_new, tz_windows=tz_windows_e),
                    "end":   graph_datetime_obj(end_local_new,   tz_windows=tz_windows_e),
                    "isReminderOn": bool(rem_minutes_for_graph_e > 0),
                    "reminderMinutesBeforeStart": int(rem_minutes_for_graph_e),
                }
                if event_type_e == "In-person":
                    patch_payload["location"] = {"displayName": location_e or ""}
                else:
                    patch_payload["location"] = {"displayName": ""}

                # Core Outlook update
                update_outlook_event(token, GRAPH["shared_mailbox_upn"], ev["outlook_event_id"], patch_payload)

                # Update ONLY the red Meeting Manager block (preserve ID & formatting)
                ok_mgr = update_outlook_manager_block(
                    outlook_event_id=ev["outlook_event_id"],
                    manager_name=manager_name_e,
                    mailbox_upn=GRAPH["shared_mailbox_upn"],
                    token=token,
                )
                if not ok_mgr:
                    st.warning("Could not update the Meeting Manager line in the Outlook body (non-fatal).")

                # 🔥 New: Update Client + Accreditation lines (insert/replace before manager block)
                ok_meta = upsert_outlook_client_and_accreditation(
                    token=token,
                    mailbox_upn=GRAPH["shared_mailbox_upn"],
                    event_id=ev["outlook_event_id"],
                    client_value=client_value,                  # from edit form
                    accreditation_required=accreditation_required_e,
                )
                if not ok_meta:
                    st.warning("Could not update Client/Accreditation lines in the Outlook body (non-fatal).")


                # ---- UPDATE Supabase ----
                supabase.table("events").update({
                    "subject": subject_e,
                    "client": (ev.get("client") or None),  # keep stored client, or extend UI to change if desired
                    "start_dt_utc": start_utc_new.isoformat(),
                    "end_dt_utc": end_utc_new.isoformat(),
                    "timezone_display": iana_new,
                    "is_all_day": is_all_day_e,
                    "location": location_e or None,
                    "event_type": ("virtual" if event_type_e == "Virtual" else "in_person"),
                    "virtual_provider": (virtual_provider_e or None),
                    "virtual_link": (virtual_link_e or None),
                    "meeting_manager_name": manager_name_e,
                    "meeting_manager_email": manager_email_e,
                    "reminder_minutes": int(rem_minutes_for_graph_e),
                    "accreditation_required": bool(accreditation_required_e),
                    "updated_at": datetime.utcnow().isoformat(),
                }).eq("id", ev["id"]).execute()

                # ---- Upsert date-certain reminder (email via app) ----
                if rem_mode_e.startswith("On date/time") and st.session_state.get("edit_reminder_datetime_local"):
                    notify_utc_e = st.session_state["edit_reminder_datetime_local"] \
                        .replace(tzinfo=ZoneInfo(iana_new)).astimezone(ZoneInfo("UTC"))
                    upsert_custom_reminder(
                        supabase_client=supabase,
                        event_id=ev["id"],
                        notify_at_utc=notify_utc_e.isoformat(),
                        to_email=manager_email_e,
                        subject_line=f"Reminder: {subject_e}",
                        body_html=f"Reminder for {subject_e} ({ev.get('client') or ''})"
                    )

                # ---- Remove missing-link reminders if we now have a link ----
                if event_type_e == "Virtual" and virtual_link_e:
                    delete_missing_link_reminders(supabase, ev["id"])

            except Exception as e:
                st.error(f"Update failed: {e}")
                st.stop()

        st.success("Event updated.")
        st.session_state.pop("edit_confirm_no_link", None)
        # Keep the edited event visible after refresh (optional UX nicety)
        try:
            # show the month of the (new) start date and include next ~30 days
            new_start_date = start_local_new.date()
            st.session_state["edit_from"] = new_start_date.replace(day=1)
            st.session_state["edit_to"]   = new_start_date + timedelta(days=30)
            # widen client filter so we don't accidentally hide it
            st.session_state["edit_client"] = "(all)"
        except Exception:
            pass        

# -----------------------------
# Export to Word (grouped by month)
# -----------------------------
st.markdown("---")


if supabase is None:
    st.info("Supabase not configured; export disabled.")
else:
    try:
        # Only current & future events (Eastern Time → UTC conversion)
        today_et = datetime.now(ZoneInfo("America/New_York")).date()
        start_floor_utc = datetime.combine(today_et, time(0, 0), tzinfo=ZoneInfo("America/New_York")).astimezone(ZoneInfo("UTC"))

        res = (
            supabase
            .table("events")
            .select("*")
            .gte("start_dt_utc", start_floor_utc.isoformat())
            .order("start_dt_utc", desc=False)
            .execute()
        )
        events = res.data or []
    except Exception as e:
        events = []
        st.error(f"Failed to load events: {e}")


    def month_key(dt: datetime) -> str:
        return dt.strftime("%B %Y").upper()

    def _fmt_hhmm(dt):
        """Portable 12-hour time like 9:00 AM (works on Windows/macOS/Linux)."""
        s = dt.strftime("%I:%M %p")  # e.g., "09:00 AM"
        return s.lstrip("0")         # -> "9:00 AM"

    def fmt_time_window_local(start_et, end_et, is_all_day, tz_label="ET"):
        if is_all_day:
            return ""  # no inline times for all-day
        start_str = _fmt_hhmm(start_et)
        end_str = _fmt_hhmm(end_et)
        return f" {start_str}–{end_str} {tz_label}"

def build_doc(events: list[dict]) -> bytes:
    doc = Document()
    title = doc.add_paragraph("Lutine Meetings Calendar")
    title_format = title.runs[0].font
    title_format.size = Pt(16)
    title_format.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    grouped = {}
    for ev in events:
        try:
            start_utc = datetime.fromisoformat(ev["start_dt_utc"].replace("Z", "+00:00"))
        except Exception:
            continue
        start_et = start_utc.astimezone(ZoneInfo("America/New_York"))
        key = month_key(start_et)
        grouped.setdefault(key, []).append((start_et, ev))

    # sort months chronologically
    for mon in sorted(grouped.keys(), key=lambda k: datetime.strptime(k, "%B %Y").date()):
        doc.add_paragraph("")  # spacing
        h = doc.add_paragraph(mon)
        h.runs[0].font.bold = True

        # iterate events within this month
        for start_et, ev in grouped[mon]:
            end_utc = datetime.fromisoformat(ev["end_dt_utc"].replace("Z", "+00:00"))
            end_et = end_utc.astimezone(ZoneInfo("America/New_York"))
            is_all_day = bool(ev.get("is_all_day"))

            # Line 1: Date/time + Client (client here; not in subject)
            month_abbr = start_et.strftime("%b")
            day_num = start_et.day
            time_win = fmt_time_window_local(start_et, end_et, is_all_day, tz_label="ET")  # e.g., " 9:00–10:00 ET"
            line1 = f"{month_abbr} {day_num}:{time_win}"
            p1 = doc.add_paragraph(line1)
            client_txt = ev.get("client") or ""
            if client_txt:
                r_client = p1.add_run(f" • Client: {client_txt}")
                # (kept in normal styling)

            # Line 2: Subject (+ location/virtual)
            subject_core = ev.get("subject") or "(No subject)"
            loc_or_v = ""
            if ev.get("event_type") == "in_person" and ev.get("location"):
                loc_or_v = ev["location"]
            elif ev.get("event_type") == "virtual":
                vp = (ev.get("virtual_provider") or "other").lower()
                loc_or_v = {"teams": "Teams", "zoom": "Zoom"}.get(vp, "Virtual")

            subject_display = subject_core
            if loc_or_v:
                subject_display += f" ({loc_or_v})"

            p2 = doc.add_paragraph(subject_display)

            # Line 3: Meeting Manager (red + bold) + Accreditation flag
            manager = ev.get("meeting_manager_name") or ""
            acc = "Y" if ev.get("accreditation_required") else "N"

            p3 = doc.add_paragraph()
            if manager:
                r_mm_label = p3.add_run("Meeting Manager: ")
                r_mm_label.bold = True
                r_mm_label.font.color.rgb = RGBColor(192, 0, 0)  # red

                r_mm_name = p3.add_run(manager)
                r_mm_name.bold = True
                r_mm_name.font.color.rgb = RGBColor(192, 0, 0)  # red

                p3.add_run("   ")  # small spacer

            r_acc_label = p3.add_run("Accreditation: ")
            r_acc_label.bold = True
            p3.add_run(acc)

            # Optional: add a small spacer line between events
            # doc.add_paragraph("")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


    if st.button("Build Word Document"):
        if not events:
            st.warning("No events found to export.")
        else:
            doc_bytes = build_doc(events)
            st.download_button(
                "Download Word (DOCX)",
                data=doc_bytes,
                file_name=f"Lutine_Master_Calendar_{date.today().year}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

# ----- Admin panel: Export to Word (always visible) -----
with st.sidebar.expander("Admin: Export Events to Word", expanded=False):
    # Filters for export (independent of page tabs)
    default_from = date.today().replace(day=1)
    default_to   = date.today() + timedelta(days=120)

    ex_from = st.date_input("From (ET)", value=st.session_state.get("export_from", default_from), key="export_from")
    ex_to   = st.date_input("To (ET)",   value=st.session_state.get("export_to", default_to),   key="export_to")

    # Optional client filter
    clients = load_clients()
    ex_client = st.selectbox("Client (optional)", ["(all)"] + clients,
                             index=st.session_state.get("export_client_idx", 0), key="export_client")

    # Cache the export so repeat downloads are instant
    @st.cache_data(ttl=300)
    def _load_events_and_build_doc(from_d: date, to_d: date, client_filter: str) -> bytes:
        # Fetch events (ET → UTC)
        start_floor_utc = datetime.combine(from_d, time(0, 0), tzinfo=ZoneInfo("America/New_York")).astimezone(ZoneInfo("UTC"))
        end_ceil_utc    = datetime.combine(to_d,   time(23,59), tzinfo=ZoneInfo("America/New_York")).astimezone(ZoneInfo("UTC"))

        q = supabase.table("events").select("*") \
                 .gte("start_dt_utc", start_floor_utc.isoformat()) \
                 .lte("start_dt_utc", end_ceil_utc.isoformat()) \
                 .order("start_dt_utc", desc=False)
        if client_filter and client_filter != "(all)":
            q = q.eq("client", client_filter)
        events = (q.execute().data or [])

        return build_doc(events)  # uses your existing build_doc(...)

    if st.button("Build Word"):
        try:
            doc_bytes = _load_events_and_build_doc(ex_from, ex_to, ex_client)
            st.session_state["export_doc_bytes"] = doc_bytes
            st.success("Export ready below.")
        except Exception as e:
            st.error(f"Export failed: {e}")

    if "export_doc_bytes" in st.session_state:
        st.download_button(
            "Download Word (DOCX)",
            data=st.session_state["export_doc_bytes"],
            file_name=f"Lutine_Master_Calendar_{date.today().year}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="export_download_btn"
        )
# -----------------------------
# Admin Tools (Sidebar)
# -----------------------------
with st.sidebar:
    st.header("Admin")

    with st.expander("Admin Tools", expanded=False):
        st.caption("Calendar sync and maintenance")

        # --- Bulk Sync (delta) ---
        if st.button("🔄 Bulk Sync Now (Outlook → App)", key="sb_bulk_sync"):
            try:
                if not GRAPH or not all(GRAPH.get(k) for k in ("tenant_id","client_id","client_secret","shared_mailbox_upn")):
                    raise RuntimeError("Missing Graph secrets.")
                if supabase is None:
                    raise RuntimeError("Supabase not configured.")

                token = get_graph_token(GRAPH["tenant_id"], GRAPH["client_id"], GRAPH["client_secret"])
                dlink = get_delta_link()  # None first time

                # First-time window: last 180d to next 365d
                if not dlink:
                    from datetime import datetime, timedelta, timezone
                    start_iso = (datetime.now(timezone.utc) - timedelta(days=180)).isoformat()
                    end_iso   = (datetime.now(timezone.utc) + timedelta(days=365)).isoformat()
                else:
                    start_iso = end_iso = None

                total_updates = 0
                last_delta = None
                for page in graph_delta_events(token, GRAPH["shared_mailbox_upn"], start_iso, end_iso, delta_link=dlink):
                    values = page.get("value", [])
                    for g in values:
                        # Skip deletions in delta
                        if "@removed" in g:
                            # Optional: also delete locally by outlook_event_id here if desired
                            continue
                        oeid = g.get("id")
                        if not oeid:
                            continue
                        res = supabase.table("events").select("id").eq("outlook_event_id", oeid).limit(1).execute()
                        rows = res.data or []
                        if not rows:
                            # Not created by app -> ignore (or insert if you want to ingest external events)
                            continue
                        row_id = rows[0]["id"]
                        updates = map_graph_event_to_row_updates(g)
                        if updates:
                            updates["updated_at"] = datetime.utcnow().isoformat()
                            supabase.table("events").update(updates).eq("id", row_id).execute()
                            total_updates += 1
                    last_delta = page.get("@odata.deltaLink") or last_delta

                if last_delta:
                    save_delta_link(last_delta)
                st.success(f"Bulk sync complete. Updated {total_updates} event(s).")
            except Exception as e:
                st.error(f"Bulk sync failed: {e}")

        st.divider()

        # --- Per-event Refresh ---
        # Pre-fill with the selected event's ID if available
        prefill_id = (ev.get("outlook_event_id") or "")
        selected_event_id = st.text_input("Outlook Event ID", value=prefill_id, key="sb_refresh_id")

        if st.button("🔃 Refresh Selected Event", key="sb_refresh_btn"):
            ev_id_raw = (selected_event_id or "").strip()
            if not ev_id_raw:
                st.warning("Enter an Outlook Event ID first.")
            else:
                try:
                    # quick sanity: Graph IDs are usually long-ish; obvious truncation -> warn early
                    if len(ev_id_raw) < 40:
                        st.warning("That ID looks truncated. Please paste the full Outlook event ID.")
                        st.stop()

                    if not GRAPH or not all(GRAPH.get(k) for k in ("tenant_id","client_id","client_secret","shared_mailbox_upn")):
                        raise RuntimeError("Missing Graph secrets.")
                    if supabase is None:
                        raise RuntimeError("Supabase not configured.")

                    token = get_graph_token(GRAPH["tenant_id"], GRAPH["client_id"], GRAPH["client_secret"])
                    g = graph_get_event(token, GRAPH["shared_mailbox_upn"], ev_id_raw)  # encodes internally

                    # Try exact match first
                    res = supabase.table("events").select("id").eq("outlook_event_id", ev_id_raw).limit(1).execute()
                    rows = res.data or []
                    if not rows and ev.get("id"):
                        # fall back: if user typed a different ID but we have a row selected, use it
                        rows = [{"id": ev["id"]}]

                    if not rows:
                        st.warning("No local event matches this Outlook ID.")
                    else:
                        row_id = rows[0]["id"]
                        updates = map_graph_event_to_row_updates(g)
                        if not updates:
                            st.info("No Outlook-owned fields to update.")
                        else:
                            updates["updated_at"] = datetime.utcnow().isoformat()
                            supabase.table("events").update(updates).eq("id", row_id).execute()
                            st.success(f"Refreshed from Outlook → updated: {', '.join(updates.keys())}")
                except Exception as e:
                    st.error(f"Refresh failed: {e}")


        st.divider()

        # (Optional) Quick list of recent events for copy/paste of IDs
        if st.checkbox("Show recent events (IDs)", key="sb_show_ids"):
            try:
                res = (
                    supabase.table("events")
                    .select("subject,start_dt_utc,outlook_event_id")
                    .order("start_dt_utc", desc=True)
                    .limit(20)
                    .execute()
                )
                rows = res.data or []
                if not rows:
                    st.caption("No events found.")
                else:
                    for r in rows:
                        st.write(f"• {r['subject']} — {r['start_dt_utc'][:16]}  |  ID: {r.get('outlook_event_id') or '—'}")
            except Exception as e:
                st.warning(f"Could not load events: {e}")



# -----------------------------
# Sidebar help
# -----------------------------
#with st.sidebar:
    #st.header("Setup Checklist")
    #st.markdown("- Shared mailbox set (e.g., calendar@yourorg.org)")
    #st.markdown("- Azure App Registration with **Calendars.ReadWrite** (Application)")
    #st.markdown("- Streamlit secrets: **graph**, **supabase**")
    #st.markdown("- Optional SMTP secrets for email: **smtp** (host, port, user, password, from_addr, from_name)")
    #st.caption("Time zones: stored as UTC + IANA; Graph uses Windows TZ IDs. Events are created with showAs=Free. Accreditation email sent if selected.")


